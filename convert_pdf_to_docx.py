"""Convert PDF to DOCX for backend processing."""
from pathlib import Path
import PyPDF2
from docx import Document
from docx.shared import Pt

pdf_path = r"C:\Users\horod\Desktop\CV_Mariusz_Horodecki_Dec_Group_Dietrich_engineering_consultan_2026-02-04.pdf"
docx_path = r"C:\Users\horod\Desktop\CV_Mariusz_Horodecki_Dec_Group_Dietrich_engineering_consultan_2026-02-04.docx"

pdf_file = Path(pdf_path)
if not pdf_file.exists():
    print(f"❌ Plik nie znaleziony: {pdf_path}")
    exit(1)

print(f"📄 Ekstrakcja tekstu z PDF ({pdf_file.stat().st_size / 1024:.1f} KB)...")
try:
    # Extract text from PDF
    text_content = []
    with open(pdf_file, 'rb') as pdf:
        reader = PyPDF2.PdfReader(pdf)
        print(f"   Liczba stron: {len(reader.pages)}")
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                text_content.append(text)
                print(f"   Strona {i+1}: {len(text)} znaków")

    if not text_content:
        print(f"⚠️  PDF nie zawiera tekstu (być może wymaga OCR)")
        exit(1)

    # Create DOCX from extracted text
    doc = Document()
    for page_text in text_content:
        for line in page_text.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.style = 'Normal'
        doc.add_paragraph()  # Add space between pages

    doc.save(str(docx_path))
    docx_file = Path(docx_path)
    print(f"\n✅ DOCX utworzony: {docx_path}")
    print(f"   Rozmiar: {docx_file.stat().st_size / 1024:.1f} KB")
except Exception as e:
    print(f"❌ Błąd konwersji: {e}")
    import traceback
    traceback.print_exc()
    exit(1)
