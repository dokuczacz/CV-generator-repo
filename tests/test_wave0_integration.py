#!/usr/bin/env python3
"""
Wave 0 Integration Tests - Against Local Azure Functions Endpoint
Tests: Idempotency (0.1), FSM transitions (0.2), Single-call execution (0.3)

Run with: python tests/test_wave0_integration.py
Requires: Azure Functions emulator running on port 7071
"""

import json
import base64
import time
from pathlib import Path
import requests
from datetime import datetime
import sys
import os

BASE_URL = "http://localhost:7071/api"

# ============================================================================
# LOGGING
# ============================================================================

def log_header(msg):
    print("\n" + "="*70)
    print(msg.center(70))
    print("="*70 + "\n")

def log_step(msg):
    print(f"[*] {msg}")

def log_ok(msg):
    print(f"[OK] {msg}")

def log_warn(msg):
    print(f"[!] {msg}")

def log_error(msg):
    print(f"[ERROR] {msg}")

def log_data(msg):
    print(f"    {msg}")

# ============================================================================
# PRE-FLIGHT
# ============================================================================

def test_health_check():
    """Verify local function is running"""
    log_header("PRE-FLIGHT: Health Check")
    
    log_step("Testing connectivity to local Azure Functions endpoint...")
    try:
        payload = {"tool_name": "cleanup_expired_sessions", "params": {}}
        resp = requests.post(f"{BASE_URL}/cv-tool-call-handler", json=payload, timeout=5)
        if resp.status_code == 200:
            log_ok("Function endpoint is ready")
            return True
        else:
            log_warn(f"Endpoint returned: {resp.status_code}")
            return False
    except Exception as e:
        log_error(f"Cannot reach endpoint: {e}")
        log_error("Run: func start")
        return False

# ============================================================================
# TEST 1: Session Creation
# ============================================================================

def get_sample_docx_base64():
    """Create or find sample DOCX"""
    try:
        sys.path.insert(0, str(Path(__file__).parent.parent))
        from tests.aline_keller_cv_data import CV_DATA
        from docx import Document
        import io
        
        log_step("Generating test DOCX from CV data...")
        doc = Document()
        
        # Add some CV content
        if isinstance(CV_DATA, dict):
            doc.add_heading("Curriculum Vitae", level=1)
            if "contact" in CV_DATA:
                doc.add_paragraph(f"Email: {CV_DATA.get('contact', {}).get('email', 'test@example.com')}")
            if "summary" in CV_DATA:
                doc.add_paragraph(str(CV_DATA["summary"])[:100])
        else:
            doc.add_heading("Test CV", level=1)
            doc.add_paragraph("This is a test CV for Wave 0 integration testing.")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return base64.b64encode(buffer.getvalue()).decode('ascii')
    except Exception as e:
        log_error(f"Could not generate test DOCX: {e}")
        return None

def test_create_session():
    """TEST 1: Create CV session"""
    log_header("TEST 1: Create CV Session")
    
    docx_b64 = get_sample_docx_base64()
    if not docx_b64:
        log_error("Could not create test DOCX")
        return None
    
    payload = {
        "tool_name": "extract_and_store_cv",
        "params": {
            "docx_base64": docx_b64,
            "language": "en",
            "extract_photo": False
        }
    }
    
    log_step("POST extract_and_store_cv...")
    try:
        resp = requests.post(f"{BASE_URL}/cv-tool-call-handler", json=payload, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        
        session_id = data.get("session_id")
        if not session_id:
            log_error(f"No session_id in response")
            log_data(f"Response: {json.dumps(data, indent=2)}")
            return None
        
        log_ok(f"Session created: {session_id}")
        log_data(f"Summary: {data.get('cv_data_summary', 'N/A')[:50]}...")
        return session_id
    except Exception as e:
        log_error(f"extract_and_store_cv failed: {e}")
        return None

# ============================================================================
# TEST 2: Baseline State
# ============================================================================

def test_baseline_state(session_id):
    """TEST 2: Get baseline session state"""
    log_header("TEST 2: Baseline Session State")
    
    log_step("GET session state...")
    payload = {
        "tool_name": "get_cv_session",
        "session_id": session_id,
        "params": {}
    }
    
    try:
        resp = requests.post(f"{BASE_URL}/cv-tool-call-handler", json=payload, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        
        stage = data.get("metadata", {}).get("stage", "UNKNOWN")
        readiness = data.get("readiness", {})
        
        log_ok(f"Current stage: {stage}")
        log_data(f"can_generate: {readiness.get('can_generate', 'N/A')}")
        log_data(f"confirmation_required: {readiness.get('confirmation_required', 'N/A')}")
        
        return stage
    except Exception as e:
        log_error(f"get_cv_session failed: {e}")
        return None

# ============================================================================
# TEST 3: Wave 0.1 - Idempotency Latch
# ============================================================================

def test_idempotency_first_gen(session_id):
    """TEST 3a: First PDF generation"""
    log_header("TEST 3a (Wave 0.1): First PDF Generation")
    
    # Confirm fields
    log_step("Confirming contact fields...")
    confirm_payload = {
        "tool_name": "update_cv_field",
        "session_id": session_id,
        "params": {
            "confirm": {
                "contact_confirmed": True,
                "education_confirmed": True
            }
        }
    }
    
    try:
        resp = requests.post(f"{BASE_URL}/cv-tool-call-handler", json=confirm_payload, timeout=10)
        resp.raise_for_status()
        log_ok("Fields confirmed")
    except Exception as e:
        log_error(f"Confirmation failed: {e}")
        return None
    
    # Generate PDF
    log_step("Requesting PDF generation (first call)...")
    gen_payload = {
        "tool_name": "process_cv_orchestrated",
        "params": {
            "message": "generate pdf",
            "session_id": session_id,
            "language": "en"
        }
    }
    
    try:
        t0 = time.time()
        resp = requests.post(f"{BASE_URL}/cv-tool-call-handler", json=gen_payload, timeout=60)
        elapsed = time.time() - t0
        resp.raise_for_status()
        data = resp.json()
        
        pdf_b64 = data.get("pdf_base64", "")
        pdf_size = len(base64.b64decode(pdf_b64)) if pdf_b64 else 0
        stage = data.get("stage", "UNKNOWN")
        
        log_ok(f"PDF generated: {pdf_size} bytes, stage={stage}, time={elapsed:.2f}s")
        
        return {
            "pdf_b64": pdf_b64,
            "pdf_size": pdf_size,
            "stage": stage,
            "time": elapsed
        }
    except Exception as e:
        log_error(f"PDF generation failed: {e}")
        return None

def test_idempotency_second_gen(session_id, first_gen):
    """TEST 3b: Second PDF generation (should use latch)"""
    log_header("TEST 3b (Wave 0.1): Second PDF Generation (Latch Test)")
    
    log_step("Requesting PDF generation (second call, should use cache)...")
    gen_payload = {
        "tool_name": "process_cv_orchestrated",
        "params": {
            "message": "generate pdf",
            "session_id": session_id,
            "language": "en"
        }
    }
    
    try:
        t0 = time.time()
        resp = requests.post(f"{BASE_URL}/cv-tool-call-handler", json=gen_payload, timeout=60)
        elapsed = time.time() - t0
        resp.raise_for_status()
        data = resp.json()
        
        pdf_b64 = data.get("pdf_base64", "")
        pdf_size = len(base64.b64decode(pdf_b64)) if pdf_b64 else 0
        stage = data.get("stage", "UNKNOWN")
        
        log_ok(f"PDF returned: {pdf_size} bytes, stage={stage}, time={elapsed:.2f}s")
        
        # Check idempotency
        if pdf_b64 and first_gen and pdf_b64 == first_gen["pdf_b64"]:
            log_ok("[PASS] Idempotency latch working - same PDF returned")
            return True
        elif pdf_size == 0:
            log_warn("[FAIL] PDF is empty (0 bytes)")
            return False
        else:
            log_warn("[FAIL] Different PDF returned - latch may not have engaged")
            return False
    except Exception as e:
        log_error(f"Second generation failed: {e}")
        return False

# ============================================================================
# TEST 4: Wave 0.2 - FSM Edit Intent
# ============================================================================

def test_fsm_edit_intent(session_id):
    """TEST 4 (Wave 0.2): FSM transitions with edit intent"""
    log_header("TEST 4 (Wave 0.2): FSM Edit Intent")
    
    log_step("Sending edit intent to escape DONE state...")
    edit_payload = {
        "tool_name": "process_cv_orchestrated",
        "params": {
            "message": "change work experience",
            "session_id": session_id,
            "language": "en"
        }
    }
    
    try:
        resp = requests.post(f"{BASE_URL}/cv-tool-call-handler", json=edit_payload, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        
        stage = data.get("stage", "UNKNOWN")
        run_summary = data.get("run_summary", {})
        stage_debug = run_summary.get("stage_debug", {})
        
        next_stage = stage_debug.get("next_stage", "UNKNOWN")
        edit_detected = stage_debug.get("edit_intent", False)
        
        log_ok(f"Edit intent processed: stage={stage}, next_stage={next_stage}")
        log_data(f"Edit intent detected: {edit_detected}")
        
        # Verify FSM transition
        if next_stage == "REVIEW":
            log_ok("[PASS] FSM transitioned to REVIEW (edit intent escape working)")
            return True
        else:
            log_warn(f"[FAIL] FSM transitioned to {next_stage} (expected REVIEW)")
            return False
    except Exception as e:
        log_error(f"Edit intent test failed: {e}")
        return False

# ============================================================================
# TEST 5: Wave 0.3 - Single-Call Execution
# ============================================================================

def test_single_call_execution(session_id):
    """TEST 5 (Wave 0.3): Single OpenAI call limit in execute stage"""
    log_header("TEST 5 (Wave 0.3): Single-Call Execution Contract")
    
    log_step("Requesting PDF generation to check OpenAI call count...")
    gen_payload = {
        "tool_name": "process_cv_orchestrated",
        "params": {
            "message": "generate pdf",
            "session_id": session_id,
            "language": "en"
        }
    }
    
    try:
        resp = requests.post(f"{BASE_URL}/cv-tool-call-handler", json=gen_payload, timeout=60)
        resp.raise_for_status()
        data = resp.json()
        
        run_summary = data.get("run_summary", {})
        model_calls = run_summary.get("model_calls", -1)
        max_calls = run_summary.get("max_model_calls", -1)
        execution_mode = run_summary.get("execution_mode", False)
        
        log_ok(f"Model calls: {model_calls}, max_limit: {max_calls}")
        log_data(f"Execution mode: {execution_mode}")
        
        # Verify single-call contract
        if execution_mode and model_calls == 1:
            log_ok("[PASS] Single-call execution working (execution_mode=True, calls=1)")
            return True
        elif execution_mode and model_calls > 1:
            log_warn(f"[FAIL] Single-call contract broken (execution_mode=True but calls={model_calls})")
            return False
        else:
            log_warn(f"[WARN] execution_mode={execution_mode}, calls={model_calls} (expected single-call in execute phase)")
            return False
    except Exception as e:
        log_error(f"Single-call test failed: {e}")
        return False

# ============================================================================
# MAIN
# ============================================================================

def main():
    log_header("Wave 0 Integration Tests - Local Azure Functions")
    log_step(f"Endpoint: {BASE_URL}")
    log_step(f"Timestamp: {datetime.now().isoformat()}")
    
    # Pre-flight
    if not test_health_check():
        log_error("Aborting: endpoint not reachable")
        return 1
    
    # Test 1: Create session
    session_id = test_create_session()
    if not session_id:
        log_error("Aborting: could not create session")
        return 1
    
    results = {}
    
    # Test 2: Baseline
    stage = test_baseline_state(session_id)
    
    # Test 3: Idempotency latch
    first_gen = test_idempotency_first_gen(session_id)
    if first_gen:
        results["0.1_latch"] = test_idempotency_second_gen(session_id, first_gen)
    else:
        results["0.1_latch"] = False
    
    # Test 4: FSM edit intent
    results["0.2_fsm"] = test_fsm_edit_intent(session_id)
    
    # Test 5: Single-call execution
    results["0.3_single_call"] = test_single_call_execution(session_id)
    
    # Summary
    log_header("Test Summary")
    for test_name, passed in results.items():
        status = "[PASS]" if passed else "[FAIL]"
        log_ok(f"{status} {test_name}")
    
    log_ok(f"Session: {session_id}")
    
    passed_count = sum(1 for v in results.values() if v)
    total = len(results)
    log_ok(f"Result: {passed_count}/{total} tests passed")
    
    return 0 if passed_count == total else 1

if __name__ == "__main__":
    sys.exit(main())
